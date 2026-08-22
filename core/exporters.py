# === exporters.py v4 (formatting: tight header, justified paras, ➤ bullets) ===
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def _contact_line(res):
    c = res.get("contact") or {}
    return " | ".join(x for x in (c.get("location"), c.get("phone"), c.get("email")) if x)


def _ascii(t):
    t = ((t or "")
         .replace("•", "-").replace("–", "-").replace("—", "-")
         .replace("’", "'").replace("‘", "'")
         .replace("“", '"').replace("”", '"')
         .replace("…", "...").replace("→", "->").replace("·", "-")
         .replace("\u00a0", " "))
    return t.encode("latin-1", "replace").decode("latin-1")

def _fmt_score(sc):
    """Normalize score display: '8.29' -> 'CGPA 8.29', '90.2' -> '90.2%'."""
    sc = str(sc or "").strip()
    if not sc:
        return ""
    low = sc.lower()
    m = re.search(r"[\d.]+", sc)
    val = m.group(0) if m else ""
    if "%" in sc:
        return sc if sc.endswith("%") else sc + "%"
    if "cgpa" in low or "gpa" in low:
        return f"CGPA {val}" if val else sc
    if "percent" in low or "marks" in low or "grade" in low:
        return f"{val}%" if val else sc
    if val:  # bare number heuristic: <=10 => CGPA, else percentage
        try:
            return f"CGPA {val}" if float(val) <= 10 else f"{val}%"
        except Exception:
            return sc
    return sc


def build_docx(res):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # justify ALL paragraphs by default

    def tight(p):
        """Header paragraphs: single line spacing, minimal gaps."""
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        return p

    def para(text):
        """Body paragraph: justified."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)
        return p

    def bullet(text):
        """➤ bullet with hanging indent, justified."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.add_run("➤ " + text)
        return p

    def sec(title):
        doc.add_heading(title, level=1)

    # ---- header (centered, line spacing 1) ----
    c = res.get("contact") or {}
    tight(doc.add_heading(c.get("name") or "Candidate", level=0)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    line = _contact_line(res)
    if line:
        tight(doc.add_paragraph(line)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if res.get("headline"):
        p = tight(doc.add_paragraph(res["headline"]))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    # ---- body (justified, ➤ bullets) ----
    if res.get("summary"):
        sec("Professional Summary")
        para(res["summary"])
    if res.get("skills"):
        sec("Core Competencies & Skills")
        para(", ".join(res["skills"]))
    if res.get("experience"):
        sec("Professional Experience")
        for exp in res["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')} | {exp.get('company', '')} | {exp.get('dates', '')}").bold = True
            for b in exp.get("bullets") or []:
                bullet(b)
    if res.get("projects"):
        sec("Key Projects")
        for proj in res["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            for b in proj.get("bullets") or []:
                bullet(b)
    if res.get("education"):
        sec("Education")
        for ed in res["education"]:
            sc = _fmt_score(ed.get("score"))
            p = doc.add_paragraph()
            p.add_run(str(ed.get("degree") or "") +
                      (f" — {ed.get('institution') or ''}" if ed.get("institution") else "")).bold = True
            line2 = "  |  ".join(x for x in (sc, str(ed.get("year") or "")) if x)
            if line2:
                q = doc.add_paragraph(line2)
                q.paragraph_format.left_indent = Inches(0.25)
    if res.get("certifications"):
        sec("Certifications")
        for x in res["certifications"]:
            bullet(x)
    if res.get("languages"):
        sec("Languages")
        para(", ".join(res["languages"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(res):
    pdf = FPDF(unit="mm", format="A4")
    pdf.set_margins(10, 10, 10)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    c = res.get("contact") or {}
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, _ascii(c.get("name") or "Candidate"),
                   align="C", new_x="LMARGIN", new_y="NEXT")
    line = _ascii(_contact_line(res))
    if line:
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(0, 5, line, align="C", new_x="LMARGIN", new_y="NEXT")
    if res.get("headline"):
        pdf.set_font("Helvetica", "I", 10)
        pdf.multi_cell(0, 6, _ascii(res["headline"]),
                       align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    def sec(title):
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, _ascii(title).upper(), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)

    def para(text):
        pdf.multi_cell(0, 5, _ascii(text), new_x="LMARGIN", new_y="NEXT", align="J")
        pdf.ln(1)

    def bullets(items):
        for b in items:
            pdf.multi_cell(0, 5, "- " + _ascii(b), new_x="LMARGIN", new_y="NEXT", align="J")
        pdf.ln(1)

    if res.get("summary"):
        sec("Professional Summary")
        para(res["summary"])
    if res.get("skills"):
        sec("Core Competencies & Skills")
        para(", ".join(res["skills"]))
    if res.get("experience"):
        sec("Professional Experience")
        for exp in res["experience"]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, _ascii(f"{exp.get('title', '')} | {exp.get('company', '')} | {exp.get('dates', '')}"),
                     new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            bullets(exp.get("bullets") or [])
    if res.get("projects"):
        sec("Key Projects")
        for proj in res["projects"]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, _ascii(proj.get("name", "")), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            bullets(proj.get("bullets") or [])
    if res.get("education"):
        sec("Education")
        for ed in res["education"]:
            sc = _fmt_score(ed.get("score"))
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 6, _ascii(str(ed.get("degree") or "") +
                           (f" - {ed.get('institution') or ''}" if ed.get("institution") else "")),
                           new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            line2 = "  |  ".join(x for x in (sc, str(ed.get("year") or "")) if x)
            if line2:
                pdf.set_x(pdf.l_margin + 3)
                pdf.multi_cell(0, 5, _ascii(line2), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
    if res.get("certifications"):
        sec("Certifications")
        bullets(res["certifications"])
    if res.get("languages"):
        sec("Languages")
        para(", ".join(res["languages"]))

    return bytes(pdf.output())